import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set custom margins (padding) for a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_horizontal_line(doc, color_hex="CCCCCC", size=12):
    """Add a thin horizontal separator line below a paragraph in python-docx."""
    p = doc.paragraphs[-1]
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{size}" w:space="4" w:color="{color_hex}"/></w:pBdr>')
    pPr.append(pBdr)

def style_paragraph(p, before=0, after=4, line_spacing=1.1):
    """Quick helper to style paragraph spacing."""
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing

def main():
    BASE_DIR = Path(__file__).parent.parent.resolve()
    output_docx_path = BASE_DIR / "resumes" / "Mayur_Resume.docx"
    print(f"Creating Word document at: {output_docx_path}")

    doc = Document()

    # --- Page Setup (Elegant Margins for Fresher Style Page Balance) ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.page_width = Inches(8.27)  # A4 Width
        section.page_height = Inches(11.69)  # A4 Height

    # Define Theme Colors
    PRIMARY_COLOR = RGBColor(30, 41, 59)      # Slate 800 (Dark grey/blue)
    SECONDARY_COLOR = RGBColor(71, 85, 105)   # Slate 600
    TEXT_COLOR = RGBColor(15, 23, 42)         # Slate 900 (Rich black)

    # --- Configure Normal Style Font ---
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = TEXT_COLOR

    # --- 1. Header Block (2-Column Table for Profile & Photo Placeholder) ---
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False

    # Width allocation (adjusted for wider margins): Left ~ 5.1 in, Right ~ 1.77 in
    table.columns[0].width = Inches(5.1)
    table.columns[1].width = Inches(1.77)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    set_cell_margins(left_cell, top=0, bottom=0, left=0, right=0)
    set_cell_margins(right_cell, top=0, bottom=0, left=0, right=0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Populate Left Column (Header Info)
    p_name = left_cell.paragraphs[0]
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("MAYUR VIVEK REDDY")
    run_name.font.size = Pt(24)
    run_name.font.bold = True
    run_name.font.color.rgb = PRIMARY_COLOR

    p_sub = left_cell.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(8)
    run_sub = p_sub.add_run("Data Annotation & Labeling Specialist | Operations")
    run_sub.font.size = Pt(11.5)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    p_contact = left_cell.add_paragraph()
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(0)
    
    r_loc = p_contact.add_run("Latur, Maharashtra, India   •   ")
    r_phone = p_contact.add_run("+91 9960235754   •   ")
    r_email = p_contact.add_run("mayurreddy286@gmail.com")
    
    for r in [r_loc, r_phone, r_email]:
        r.font.size = Pt(10)
        r.font.color.rgb = SECONDARY_COLOR

    # Populate Right Column (Passport-sized Photo Placeholder Box)
    p_photo = right_cell.paragraphs[0]
    p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_photo.paragraph_format.space_before = Pt(0)
    p_photo.paragraph_format.space_after = Pt(0)

    # Insert a 1x1 table to act as the perfect border outline for photo placeholder
    photo_box_table = right_cell.add_table(rows=1, cols=1)
    photo_box_table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    photo_box_table.autofit = False
    
    photo_cell = photo_box_table.cell(0, 0)
    photo_cell.width = Inches(1.1)
    
    # Set height of row XML
    trPr = photo_box_table.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '1872') # 1.3 inches in twentieths of a pt
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

    # Set background shading and dashed/thin borders
    tcPr = photo_cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="dashed" w:sz="6" w:space="0" w:color="A0A0A0"/><w:left w:val="dashed" w:sz="6" w:space="0" w:color="A0A0A0"/><w:bottom w:val="dashed" w:sz="6" w:space="0" w:color="A0A0A0"/><w:right w:val="dashed" w:sz="6" w:space="0" w:color="A0A0A0"/></w:tcBorders>')
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
    tcPr.append(tcBorders)
    tcPr.append(shd)

    photo_para = photo_cell.paragraphs[0]
    photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_para.paragraph_format.space_before = Pt(25) # Centered vertically
    photo_para.paragraph_format.space_after = Pt(0)
    
    r_photo_text = photo_para.add_run("Add Photo\nHere")
    r_photo_text.font.size = Pt(8.5)
    r_photo_text.font.bold = True
    r_photo_text.font.color.rgb = SECONDARY_COLOR

    # --- Add spacing before content ---
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(12)
    p_spacer.paragraph_format.space_after = Pt(0)

    # --- Reusable Section Builder ---
    def add_section_header(title):
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(14)
        p_hdr.paragraph_format.space_after = Pt(4)
        p_hdr.paragraph_format.keep_with_next = True
        run_hdr = p_hdr.add_run(title.upper())
        run_hdr.font.size = Pt(12.5)
        run_hdr.font.bold = True
        run_hdr.font.color.rgb = PRIMARY_COLOR
        add_horizontal_line(doc, color_hex="CBD5E1", size=8) # 1pt line

    # --- 2. Objective Section ---
    add_section_header("Objective")
    p_obj = doc.add_paragraph()
    style_paragraph(p_obj, before=4, after=6)
    p_obj.paragraph_format.line_spacing = 1.15
    run_obj = p_obj.add_run(
        "Detail-oriented and analytical professional with a strong foundation in data management, accuracy, and foundational "
        "AI/ML concepts. Seeking a Data Annotation / Data Labeling role where I can apply my high attention to detail, precision "
        "focus, and systematic workflow handling to support high-quality AI training dataset generation while continuously "
        "developing my career in the AI operations domain."
    )
    run_obj.font.size = Pt(10)

    # --- 3. Skills Section ---
    add_section_header("Skills")
    skills_data = [
        ("Data & Annotation", "Data Annotation, Image/Text Labeling, Quality Focus, Data Accuracy, Validation"),
        ("AI/ML Foundations", "Foundational AI Basics, Machine Learning Concepts, Dataset Operations"),
        ("Office & Software", "Microsoft Word, Microsoft Excel (Data Entry, Lists, Spreadsheets), Basic Computer Knowledge"),
        ("Core Strengths", "High Attention to Detail, Analytical Thinking, Problem Solving, Time Management, Active Communication")
    ]
    for category, skills in skills_data:
        p_skill = doc.add_paragraph()
        style_paragraph(p_skill, before=1, after=4)
        r_cat = p_skill.add_run(f"•  {category}: ")
        r_cat.bold = True
        r_cat.font.size = Pt(10)
        r_cat.font.color.rgb = PRIMARY_COLOR
        r_skills = p_skill.add_run(skills)
        r_skills.font.size = Pt(10)

    # --- 4. Education Section ---
    add_section_header("Education")
    education_data = [
        ("Bachelor of Commerce (B.Com)", "Rajarshi Shahu Autonomous College, Latur", "2020"),
        ("Higher Secondary Certificate (HSC)", "Maharashtra State Board, Latur", "2017"),
        ("Secondary School Certificate (SSC)", "Maharashtra State Board, Latur", "2015")
    ]
    for degree, school, year in education_data:
        p_edu = doc.add_paragraph()
        style_paragraph(p_edu, before=2, after=3)
        r_deg = p_edu.add_run(f"{degree}  •  ")
        r_deg.bold = True
        r_deg.font.size = Pt(10)
        r_deg.font.color.rgb = PRIMARY_COLOR
        
        r_sch = p_edu.add_run(school)
        r_sch.font.size = Pt(10)
        
        r_yr_tab = p_edu.add_run(f"\t{year}")
        r_yr_tab.font.size = Pt(10)
        p_edu.paragraph_format.tab_stops.add_tab_stop(Inches(6.8)) # Right-aligned date Stop (adjusted for wider margins)

    # --- 5. Achievements Section ---
    add_section_header("Achievements & Extra-Curriculars")
    achievements = [
        ("National Cadet Corps (NCC)", "Participation — developed strong discipline, leadership, team collaboration, and structured execution skills under pressure."),
        ("National Service Scheme (NSS)", "Participation — engaged in volunteer-driven social welfare projects, community outreach, and systematic event reporting.")
    ]
    for title, desc in achievements:
        p_ach = doc.add_paragraph()
        style_paragraph(p_ach, before=1, after=4)
        r_ach_t = p_ach.add_run(f"•  {title}: ")
        r_ach_t.bold = True
        r_ach_t.font.size = Pt(10)
        r_ach_t.font.color.rgb = PRIMARY_COLOR
        
        r_ach_d = p_ach.add_run(desc)
        r_ach_d.font.size = Pt(9.8)

    # --- 6. Languages Section ---
    add_section_header("Languages")
    p_lang = doc.add_paragraph()
    style_paragraph(p_lang, before=4, after=0)
    r_lang = p_lang.add_run("English (Professional)  •  Hindi (Fluent)  •  Marathi (Native)  •  Kannada (Conversational)")
    r_lang.font.size = Pt(10)
    r_lang.font.bold = True
    r_lang.font.color.rgb = SECONDARY_COLOR

    # --- Save Document ---
    doc.save(str(output_docx_path))
    print("[SUCCESS] Word Document updated successfully (experience removed)!")

if __name__ == "__main__":
    main()
